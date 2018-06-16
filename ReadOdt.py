from math import fabs

import numpy as np
from odf.opendocument import load
from odf.table import Table, TableRow, TableCell
from docx import Document

# this is the row data in the table starts on
TABLE_DATA_ROW = 3
# this is the number of sessions after which corrections were applied
SESSION_NUM = 3


# function handling reading of an odt file - puts it in a list of lists, assumes there is only one table in a file
def read_table(filez, threshold, correction_sessions):
    table = []
    doc = load(filez)
    data_begin = -1
    try:
        tab = doc.text.getElementsByType(Table)[0]
        row = len(tab.getElementsByType(TableRow))
        col = len(tab.getElementsByType(TableRow)[0].getElementsByType(TableCell))
        for i in range(0, row):
            table.append([])
            try:
                float(str(tab.getElementsByType(TableRow)[i].getElementsByType(TableCell)[0]))
                if data_begin < 0:
                    data_begin = i
            except ValueError:
                pass
            for j in range(0, col):
                string = str(tab.getElementsByType(TableRow)[i].getElementsByType(TableCell)[j])
                if "," in string:
                    string = string.replace(",", ".")
                table[i].append(string)

        table = correct(table, row, threshold, correction_sessions, data_begin)
    except IndexError:
        row = 0
        col = 0
    return [row, col, table]


# analysis of a single data list
def calculate_avg_stdev(table, mean_number):
    values = [[], [], [], []]

    if mean_number == 0 or len(table) < mean_number:
        length = len(table)
    else:
        length = mean_number

    for row in range(0, length):
        for col in range(0, 4):
            try:
                val = float(table[row][col])
                values[col].append(val)
            except ValueError:
                pass

    mean = np.around([np.mean(values[0]), np.mean(values[1]), np.mean(values[2]), np.mean(values[3])], decimals=1)
    stdev = np.around([np.std(values[0]), np.std(values[1]), np.std(values[2]), np.std(values[3])], decimals=2)

    return [mean, stdev]


# function handling reading of table in  a docx file - puts it in a list of lists, assumes there is only one table
# in a file
def read_docx(filez, threshold, correction_sessions):
    table = []
    doc = Document(filez)
    tables = doc.tables
    data_begin = -1
    for tab in tables:
        for i in range(0, len(tab.rows)):
            table.append([])
            for j in range(0, len(tab.rows[i].cells)):
                for paragraph in tab.rows[i].cells[j].paragraphs:
                    string = paragraph.text
                    if "," in string:
                        string = string.replace(",", ".")
                    table[i].append(string)
            try:
                float(table[i][0])
                if data_begin < 0:
                    data_begin = i
            except ValueError:
                pass

    return [len(table), len(table[0]), correct(table, len(table), threshold, correction_sessions, data_begin)]


def correct(table, row, threshold, correction_sessions, data_begin):
    start = data_begin + correction_sessions
    if start < row:
        for i in range(start, row):
            for j in range(6, 10):
                avr = float(table[i - 1][j])
                if fabs(avr) > threshold:
                    try:
                        val = float(table[i][j - 6])
                        table[i][j - 6] = str(np.around(val - avr, decimals=1))
                    except ValueError:
                        pass
    return table
